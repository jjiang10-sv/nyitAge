from nbformat import read
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import base64
from io import BytesIO
from PIL import Image

def add_code_block(doc, code):
    """Adds a formatted code block to the Word document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    
    # Add background shading for code blocks
    shading = parse_xml(r'<w:shd {} w:fill="EDEDED"/>'.format(nsdecls('w')))
    paragraph._element.get_or_add_pPr().append(shading)

def add_output_block(doc, output):
    """Adds a formatted output block to the Word document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(output)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    
    # Different background for output blocks
    shading = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
    paragraph._element.get_or_add_pPr().append(shading)

# Load the Jupyter Notebook
notebook_path = "stack.ipynb"
with open(notebook_path, "r", encoding="utf-8") as f:
    nb = read(f, as_version=4)

# Create a new Word document
doc = Document()

# Process each cell in the notebook
for cell in nb.cells:
    if cell.cell_type == "markdown":
        doc.add_paragraph(cell.source)  # Add markdown as normal text
    
    elif cell.cell_type == "code":
        add_code_block(doc, cell.source)  # Add code with formatting
        
        # Add outputs if they exist
        if cell.outputs:
            for output in cell.outputs:
                if "text" in output:  # Extract text output
                    add_output_block(doc, output["text"])
                elif "image/png" in output.data:  # Handle image outputs
                    img_data = base64.b64decode(output.data["image/png"])
                    img = Image.open(BytesIO(img_data))
                    img_path = "temp_output.png"
                    img.save(img_path)
                    doc.add_picture(img_path)

# Save the Word document
output_path = "stack.docx"
doc.save(output_path)

print(f"Notebook successfully converted to {output_path}")
